# src/data_ingestion.py
import os
import json
import fitz  # PyMuPDF for PDF processing
from docx import Document # python-docx for DOCX processing
import pytesseract # For OCR
from PIL import Image # For image processing with OCR
import io

# --- Configuration for Tesseract (if not in PATH) ---
# IMPORTANT: If Tesseract is not in your system's PATH, uncomment and set the path below.
# pytesseract.pytesseract.tesseract_cmd = r'/path/to/tesseract' # For Linux/macOS
# pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe' # For Windows

def load_text_from_pdf(file_path):
    """Extracts text from a PDF file."""
    text = ""
    try:
        doc = fitz.open(file_path)
        for page_num in range(len(doc)):
            page = doc.load_page(page_num)
            text += page.get_text()
            # --- Optional: OCR for images embedded within PDF pages ---
            # This section can be enabled if OCR for images inside PDFs is needed.
            # Note: This can significantly slow down PDF processing and increase memory usage.
            # image_list = page.get_images(full=True)
            # for img_index, img in enumerate(image_list):
            #     xref = img[0]
            #     base_image = doc.extract_image(xref)
            #     image_bytes = base_image["image"]
            #     try:
            #         pil_image = Image.open(io.BytesIO(image_bytes))
            #         text += "\n" + pytesseract.image_to_string(pil_image)
            #     except Exception as e:
            #         print(f"OCR Error for image in PDF {file_path}: {e}")
        doc.close()
    except Exception as e:
        print(f"Error reading PDF {file_path}: {e}")
    return text

def load_text_from_docx(file_path):
    """Extracts text from a DOCX file."""
    text = ""
    try:
        doc = Document(file_path)
        for para in doc.paragraphs:
            text += para.text + "\n"
    except Exception as e:
        print(f"Error reading DOCX {file_path}: {e}")
    return text

def load_from_json(file_path):
    """Loads data from a JSON file."""
    data = None
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
    except Exception as e:
        print(f"Error reading JSON {file_path}: {e}")
    return data

def ocr_image(image_path):
    """Performs OCR on an image file."""
    text = ""
    try:
        text = pytesseract.image_to_string(Image.open(image_path))
    except Exception as e:
        print(f"Error performing OCR on image {image_path}: {e}")
    return text

def ingest_data(data_source_path):
    """
    Ingests data from various file types in a given directory or a single file.
    Returns a list of dictionaries, where each dictionary contains 'file_name' and 'content'.
    For JSON, 'content' will be the parsed JSON data.
    For text-based files (PDF, DOCX), 'content' will be the extracted text.
    """
    ingested_data_list = []
    
    if os.path.isfile(data_source_path):
        file_paths = [data_source_path]
    elif os.path.isdir(data_source_path):
        file_paths = [os.path.join(data_source_path, f) for f in os.listdir(data_source_path) if os.path.isfile(os.path.join(data_source_path, f))]
    else:
        print(f"Error: {data_source_path} is not a valid file or directory.")
        return ingested_data_list

    for file_path in file_paths:
        file_name = os.path.basename(file_path)
        content = None
        print(f"Processing: {file_name}")

        if file_path.lower().endswith('.pdf'):
            content = load_text_from_pdf(file_path)
        elif file_path.lower().endswith('.docx'):
            content = load_text_from_docx(file_path)
        elif file_path.lower().endswith('.json'):
            # For JSON, we might want to store the structured data directly
            # or extract specific text fields depending on the schema.
            # Here, we'll store the parsed JSON.
            content = load_from_json(file_path)
        elif file_path.lower().endswith(('.png', '.jpg', '.jpeg', '.tiff', '.bmp')):
            content = ocr_image(file_path) # OCR for standalone image files
        else:
            print(f"Unsupported file type: {file_name}")
            continue
        
        if content:
            ingested_data_list.append({"file_name": file_name, "raw_content": content, "source_path": file_path})
            
    return ingested_data_list

# --- Sample Usage ---
if __name__ == "__main__":
    print("Starting Data Ingestion Module...")

    # Note: For full testing of PDF and Image OCR in this sample script,
    # either manually create the sample files (e.g., 'sample_resume.pdf', 'sample_image.png')
    # or uncomment and adapt the file creation code below.
    # Create dummy files for testing
    sample_data_dir = "sample_data_ingestion"
    os.makedirs(sample_data_dir, exist_ok=True)

    # Dummy PDF (requires manual creation or a library to generate a simple PDF)
    # For now, we'll assume a PDF exists or skip this part in auto-testing.
    # with open(os.path.join(sample_data_dir, "sample_resume.pdf"), "w") as f: 
    #     f.write("%PDF-1.4\n1 0 obj<</Type/Catalog/Pages 2 0 R>>endobj\n2 0 obj<</Type/Pages/Count 1/Kids[3 0 R]>>endobj\n3 0 obj<</Type/Page/MediaBox[0 0 612 792]/Parent 2 0 R/Resources<</Font<</F1 4 0 R>>>>/Contents 5 0 R>>endobj\n4 0 obj<</Type/Font/Subtype/Type1/BaseFont/Helvetica>>endobj\n5 0 obj<</Length 44>>stream\nBT /F1 24 Tf 100 700 Td (Hello PDF World) Tj ET\nendstream\nendobj\nxref\n0 6\n0000000000 65535 f\n0000000010 00000 n\n0000000062 00000 n\n0000000111 00000 n\n0000000212 00000 n\n0000000260 00000 n\ntrailer<</Size 6/Root 1 0 R>>\nstartxref\n300\n%%EOF")
    # print("Created dummy sample_resume.pdf")

    # Dummy DOCX
    try:
        doc = Document()
        doc.add_paragraph("This is a sample resume in DOCX format.")
        doc.add_paragraph("Skills: Python, Machine Learning, Data Analysis.")
        doc.save(os.path.join(sample_data_dir, "sample_resume.docx"))
        print("Created dummy sample_resume.docx")
    except Exception as e:
        print(f"Could not create dummy DOCX (python-docx might be missing or other issue): {e}")

    # Dummy JSON
    sample_job_posting = {
        "job_id": "123",
        "title": "Data Scientist",
        "company": "Tech Solutions Inc.",
        "description": "Looking for a skilled data scientist with experience in ML.",
        "requirements": ["Python", "SQL", "Machine Learning"],
        "location": "New York, NY"
    }
    with open(os.path.join(sample_data_dir, "sample_job.json"), 'w') as f:
        json.dump(sample_job_posting, f, indent=4)
    print("Created dummy sample_job.json")

    # Dummy Image (requires Pillow to create a simple image, or use a placeholder)
    # For now, we'll assume an image exists or skip this part in auto-testing.
    # try:
    #     img = Image.new('RGB', (200, 100), color = 'red')
    #     img.save(os.path.join(sample_data_dir, "sample_image.png"))
    #     print("Created dummy sample_image.png")
    # except ImportError:
    #     print("Pillow library not found, skipping dummy image creation.")
    # except Exception as e:
    #     print(f"Could not create dummy image: {e}")

    print(f"\n--- Ingesting all data from directory: {sample_data_dir} ---")
    all_data = ingest_data(sample_data_dir)
    for item in all_data:
        print(f"\nFile: {item['file_name']}")
        # print(f"Source: {item['source_path']}") # Optional: print source path
        if isinstance(item['raw_content'], dict): # For JSON
            print("Content (JSON):")
            print(json.dumps(item['raw_content'], indent=2))
        else: # For text
            print(f"Content (Text Preview - first 200 chars):\n{item['raw_content'][:200]}...")
        print("-" * 20)

    # Example: Ingesting a single file
    # single_file_path = os.path.join(sample_data_dir, "sample_job.json")
    # if os.path.exists(single_file_path):
    #     print(f"\n--- Ingesting single file: {single_file_path} ---")
    #     single_file_data = ingest_data(single_file_path)
    #     if single_file_data:
    #         print(f"File: {single_file_data[0]['file_name']}")
    #         print(f"Content (JSON):\n{json.dumps(single_file_data[0]['raw_content'], indent=2)}")
    # else:
    #     print(f"\nSkipping single file ingestion test as {single_file_path} does not exist.")

    print("\nData Ingestion Module execution finished.")
    # Consider cleaning up dummy files/directory after testing if needed
    # import shutil
    # shutil.rmtree(sample_data_dir)
    # print(f"Cleaned up {sample_data_dir}")